from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_run_font(run, size=10, bold=False, color=None):
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_title_page(doc, title_text, subtitle_text):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_text)
    set_run_font(run, size=28, bold=True, color='1a5276')
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle_text)
    set_run_font(run2, size=12, color='666666')
    doc.add_page_break()

def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=16, bold=True, color='1a5276')

def add_sub_title(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color='2e86c1')

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=9, bold=True)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)
    doc.add_paragraph()

def add_para(doc, text, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)

# ==================== BUILD DOC ====================

add_title_page(doc,
    '嵌入式 / BSP 实习岗位\n分析报告',
    '目标：BSP 工程师\n日期：2026-07-11\nGitHub：github.com/mahu-bit/embedded-bsp-learning')

# ---- 1 ----
add_section_title(doc, '一、三类实习岗位总览')
add_table(doc,
    ['档次', '岗位类型', '薪资', '代表企业', '能投时间'],
    [
        ['第一档', 'MCU嵌入式实习生', '120~500元/天', '小米、理想、迅音科技', '2~3个月后'],
        ['第二档', 'Linux驱动/BSP', '200~400元/天', '华大半导体、海康、乐鑫', '5~6个月后'],
        ['第三档', '芯片原厂BSP', '250~500元/天', '华为海思、NXP、TI、大疆', '8~12个月后'],
    ])

# ---- 2 ----
add_section_title(doc, '二、第一档：MCU 嵌入式实习生（首选目标）')
add_para(doc, '预计能投时间：从现在起 2~3 个月', bold=True)

add_sub_title(doc, '真实岗位')
add_table(doc,
    ['企业', '岗位', '薪资', '地点', '核心要求'],
    [
        ['小米集团', '嵌入式开发实习生', '250~400元/天', '南京', 'C/C++、Python、RTOS、CAN/LIN'],
        ['迅音科技', '嵌入式系统开发', '300~500元/天', '上海', 'C/C++、QT、STM32、RTOS'],
        ['理想汽车', 'MCU嵌入式开发', '面议', '-', 'C语言、MCU、汽车电子'],
        ['BOSS直聘\n(大量中小厂)', '单片机工程师', '150~250元/天', '全国', 'STM32/51、Keil、I2C/SPI/UART'],
    ])

add_sub_title(doc, '技能对照你的学习计划')
add_table(doc,
    ['技能', '在计划中位置', '学会时间', '状态'],
    [
        ['C语言精通', '第一阶段', '正在学', '进行中'],
        ['STM32开发', '第二阶段', '1~2月后', '⬜'],
        ['UART/I2C/SPI', '第二阶段外设', '1.5月后', '⬜'],
        ['FreeRTOS基础', '第四阶段', '3~4月后', '⬜'],
        ['看懂原理图', '第三阶段', '2月后', '⬜'],
        ['Git版本管理', '已完成', '✅', '已完成'],
    ])

# ---- 3 ----
add_section_title(doc, '三、第二档：Linux驱动/BSP实习生')
add_para(doc, '预计能投时间：从现在起 5~6 个月', bold=True)

add_sub_title(doc, '真实岗位')
add_table(doc,
    ['企业', '岗位', '薪资', '地点', '核心要求'],
    [
        ['华大半导体', 'BSP软件研发', '~9K/月', '成都', 'C/C++精通、Linux内核、ARM/RISC-V'],
        ['海康威视', '嵌入式BSP', '9K~20K/月', '杭州', 'C/C++、Linux、设备驱动、ARM'],
        ['乐鑫科技', 'AI芯片BSP', '面议', '上海', 'C语言、RTOS、驱动开发'],
        ['全志科技', '嵌入式软件', '150~200元/天', '珠海', 'C语言、Linux、ARM'],
        ['溯简科技', '嵌入式软件', '6K~21K', '深圳', 'BSP维护、Linux驱动、I2C/SPI/UART'],
    ])

add_sub_title(doc, '技能对照')
add_table(doc,
    ['技能', '在计划中位置', '学会时间'],
    [
        ['C语言精通', '第一阶段', '进行中'],
        ['Linux系统编程', '第六阶段', '4~5月后'],
        ['Linux设备驱动', '第十阶段', '6~8月后'],
        ['字符设备驱动', '第十阶段', '6~8月后'],
        ['设备树', '第十阶段', '6~8月后'],
        ['内核编译裁剪', '第九阶段', '6~8月后'],
    ])

# ---- 4 ----
add_section_title(doc, '四、第三档：芯片原厂BSP（终极目标）')
add_para(doc, '预计能投时间：从现在起 8~12 个月', bold=True)

add_sub_title(doc, '目标企业')
add_table(doc,
    ['梯队', '行业', '企业'],
    [
        ['T0', '芯片原厂', '华为海思、兆易创新、乐鑫、联发科、地平线、NXP、TI、ARM'],
        ['T0', '机器人', '大疆、海康机器人、汇川、石头科技'],
        ['T1', '消费电子', '华为、小米、OPPO、vivo、大疆、海康威视'],
        ['T1', '通信设备', '华为、新华三、锐捷、TP-Link'],
    ])

add_sub_title(doc, '全部技能清单')
add_table(doc,
    ['#', '技能', '计划位置', '状态'],
    [
        ['1', 'C语言+数据结构', '第一阶段', '进行中'],
        ['2', 'STM32全部外设', '第二阶段', '⬜'],
        ['3', 'FreeRTOS', '第四阶段', '⬜'],
        ['4', 'Linux系统编程', '第六阶段', '⬜'],
        ['5', 'Linux网络编程', '第七阶段', '⬜'],
        ['6', 'ARM裸机编程', '第八阶段', '⬜'],
        ['7', 'Uboot移植', '第九阶段', '⬜'],
        ['8', '内核编译裁剪', '第九阶段', '⬜'],
        ['9', '设备树', '第十阶段', '⬜'],
        ['10', '驱动开发', '第十阶段', '⬜'],
        ['11', 'Buildroot/Yocto', '第十一阶段', '⬜'],
    ])

# ---- 5 ----
add_section_title(doc, '五、完整时间线')
add_table(doc,
    ['时间', '学什么', '可投什么', '薪资'],
    [
        ['现在~1月后', 'C语言收尾', '先打基础', '-'],
        ['1~2月后', 'STM32外设+项目1', '开始投简历', '120~250/天'],
        ['2~3月后', 'STM32进阶+RTOS', '第一档：MCU嵌入式', '150~400/天'],
        ['4~5月后', 'Linux系统编程', 'Linux岗位', '200~350/天'],
        ['6~8月后', 'Linux驱动+内核', '第二档：BSP/驱动', '4K~9K/月'],
        ['8~12月后', '完整BSP栈', '第三档：芯片原厂', '15K~25K/月'],
    ])

# ---- 6 ----
add_section_title(doc, '六、高频技能词云')
add_table(doc,
    ['热度', '技能', '你的进度'],
    [
        ['🔴🔴🔴🔴🔴', 'C语言（绝对核心）', '学习中（指针✅，在学函数）'],
        ['🔴🔴🔴🔴', 'UART/SPI/I2C', '2月后'],
        ['🔴🔴🔴🔴', 'STM32/ARM', '1~2月后'],
        ['🔴🔴🔴', 'Linux系统编程', '4~5月后'],
        ['🔴🔴🔴', 'RTOS', '3~4月后'],
        ['🔴🔴🔴', '原理图+示波器', '2月后'],
        ['🔴🔴', 'Python/Shell', '4月后'],
        ['🔴🔴', '数据结构', 'C语言后期'],
        ['🔴', 'Git版本管理', '✅已完成'],
    ])

# ---- 7 ----
add_section_title(doc, '七、现在就能做的事')
add_table(doc,
    ['#', '事项', '状态'],
    [
        ['1', 'GitHub仓库建好', '✅'],
        ['2', 'C语言学完', '约50%'],
        ['3', '买STM32F103开发板', 'C语言快学完时'],
        ['4', '项目1：串口控制LED', '⬜'],
        ['5', '项目代码传GitHub', '⬜'],
        ['6', '写简历投实习', '⬜'],
    ])

# ---- 8 ----
add_section_title(doc, '八、面试必考10题')
add_table(doc,
    ['#', '考题', '对应模块', '状态'],
    [
        ['1', 'static关键字的作用', '函数/变量', '已学'],
        ['2', 'const的用法和位置含义', '指针', '已学'],
        ['3', '指针和数组的区别', '指针', '已学'],
        ['4', '结构体内存对齐', '结构体', '近期'],
        ['5', '大端小端', '联合体', '近期'],
        ['6', '中断的概念和流程', 'STM32', '2月后'],
        ['7', 'I2C和SPI的区别', 'STM32', '2月后'],
        ['8', '进程和线程的区别', 'Linux', '4月后'],
        ['9', 'volatile关键字', '嵌入式', 'STM32阶段'],
        ['10', '链表反转', '数据结构', 'C语言后期'],
    ])

# ---- 9 ----
add_section_title(doc, '九、投递渠道')
add_table(doc,
    ['渠道', '网址', '说明'],
    [
        ['BOSS直聘', 'zhipin.com', '岗位最多，直接和HR聊'],
        ['实习僧', 'shixiseng.com', '专注实习，质量好'],
        ['牛客网', 'nowcoder.com', '内推多，面经多'],
        ['企业官网', '各公司招聘页', '芯片原厂走官网'],
        ['学校就业网', '-', '竞争相对小'],
    ])

add_para(doc, '')
add_para(doc, '搜索关键词：嵌入式实习生 | BSP实习生 | MCU实习生 | 驱动开发实习生 | 单片机开发实习生', bold=True)

# Save
output_path = 'E:/c语言大作业/学习/嵌入式实习岗位分析.docx'
doc.save(output_path)
print('Done: ' + output_path)
